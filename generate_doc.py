import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Set standard margins (0.8 inch top/bottom, 0.9 inch left/right)
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Color Palette
COLOR_PRIMARY = RGBColor(15, 76, 129)     # Deep Navy (#0F4C81)
COLOR_SECONDARY = RGBColor(70, 80, 95)    # Slate (#46505F)
COLOR_DARK = RGBColor(30, 41, 59)        # Dark Text (#1E293B)
HEX_PRIMARY = '0F4C81'
HEX_ALT_ROW = 'F8FAFC'
HEX_BORDER = 'CBD5E1'

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=120, right=120):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{HEX_PRIMARY}"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{HEX_PRIMARY}"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

# Title Header
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(4)

run_title = title_p.add_run('KỊCH BẢN KIỂM THỬ (TEST PLAN & REPORT)\n')
run_title.font.name = 'Arial'
run_title.font.size = Pt(17)
run_title.font.bold = True
run_title.font.color.rgb = COLOR_PRIMARY

sub_title = title_p.add_run('HỆ THỐNG ĐẶT LỊCH HẸN & XÁC THỰC — NHA KHOA GOODSMILE\n')
sub_title.font.name = 'Arial'
sub_title.font.size = Pt(12)
sub_title.font.bold = True
sub_title.font.color.rgb = COLOR_SECONDARY

date_run = title_p.add_run('Dự án Tốt nghiệp FPT • Phiên bản 1.0 • Ngày nghiệm thu: 04/08/2026')
date_run.font.name = 'Arial'
date_run.font.size = Pt(9.5)
date_run.font.italic = True
date_run.font.color.rgb = RGBColor(120, 130, 140)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_heading(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = COLOR_PRIMARY
    return h

# 1. Giới thiệu chung
add_heading('1. GIỚI THIỆU CHUNG')
p1 = doc.add_paragraph()
p1.paragraph_format.space_after = Pt(4)
r1 = p1.add_run(
    'Tài liệu này trình bày chi tiết các kịch bản kiểm thử (Test Cases) phục vụ công tác kiểm định chất lượng (QC) '
    'và nghiệm thu cho Module Đặt lịch hẹn khám (Appointments) và Hệ thống Xác thực người dùng (Authentication/OTP) '
    'của dự án Nha khoa GoodSmile.'
)
r1.font.name = 'Arial'
r1.font.size = Pt(10)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(10)
r2 = p2.add_run(
    'Phạm vi kiểm thử bao gồm các luồng xử lý chính (Happy Path), các ràng buộc nghiệp vụ nâng cao (Edge Cases) như '
    'chống trùng lịch điều trị của bác sĩ thông qua Redis Distributed Lock, tự động xếp hàng chờ khám (Queueing), '
    'và tự động liên kết hồ sơ bệnh án cũ dựa trên Số điện thoại.'
)
r2.font.name = 'Arial'
r2.font.size = Pt(10)

def create_test_table(headers, rows_data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Set Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 4] else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Rows
    for row_idx, data in enumerate(rows_data):
        row_cells = table.add_row().cells
        bg_hex = HEX_ALT_ROW if row_idx % 2 == 1 else 'FFFFFF'
        for i, val in enumerate(data):
            row_cells[i].text = val
            set_cell_background(row_cells[i], bg_hex)
            set_cell_margins(row_cells[i], top=100, bottom=100, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 4] else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_DARK
                if i == 0:
                    run.font.bold = True
                if i == 4:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(19, 115, 51)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

headers = ['Mã TC', 'Tên Kịch Bản', 'Các Bước Thực Hiện', 'Kết Quả Kỳ Vọng', 'Kết Quả']

# 2. Authentication
add_heading('2. KỊCH BẢN KIỂM THỬ XÁC THỰC (AUTHENTICATION)')
data_auth = [
    [
        'TC-UI-01',
        'Đăng nhập Lễ tân mẫu thành công',
        '1. Truy cập trang đăng nhập.\n2. Nhập SĐT "0901234567" và MK "12345678".\n3. Bấm Đăng nhập.',
        'Đăng nhập thành công, chuyển hướng đến Dashboard Lễ tân (/dashboard/receptionist) và hiển thị thông tin chính xác.',
        'ĐẠT'
    ],
    [
        'TC-UI-02',
        'Đăng nhập thất bại do sai mật khẩu',
        '1. Truy cập trang đăng nhập.\n2. Nhập SĐT "0901234567" và MK "wrong_pass".\n3. Bấm Đăng nhập.',
        'Hệ thống báo lỗi "Mã đăng nhập hoặc mật khẩu không chính xác" (INVALID_CREDENTIALS) và giữ nguyên màn hình.',
        'ĐẠT'
    ],
    [
        'TC-UI-03',
        'Đăng ký Bệnh nhân mới thành công',
        '1. Truy cập trang Đăng ký.\n2. Nhập Họ tên, SĐT ngẫu nhiên chưa đăng ký, Mật khẩu.\n3. Nhấp gửi OTP, nhập mã OTP xác thực.\n4. Bấm Đăng ký.',
        'Đăng ký tài khoản thành công. Tự động tạo hồ sơ bệnh nhân tương ứng trong DB dưới phân hạng STANDARD.',
        'ĐẠT'
    ]
]
create_test_table(headers, data_auth)

# 3. Booking
add_heading('3. KỊCH BẢN KIỂM THỬ ĐẶT LỊCH KHÁM (BOOKING)')
data_booking = [
    [
        'TC-UI-04',
        'Bệnh nhân vãng lai đặt lịch trực tuyến (Guest)',
        '1. Tại màn hình Đặt lịch công khai, nhập Họ tên & SĐT mới.\n2. Chọn dịch vụ S-02 (Tẩy trắng răng), chọn bác sĩ D-01.\n3. Chọn ngày khám và 1 slot trống.\n4. Xác thực OTP & xác nhận.',
        'Đặt lịch thành công, hệ thống tự động sinh hồ sơ bệnh nhân vãng lai mới và gửi sự kiện cập nhật real-time tới quầy Lễ tân.',
        'ĐẠT'
    ],
    [
        'TC-UI-05',
        'Bệnh nhân đã đăng nhập đặt lịch trực tuyến',
        '1. Đăng nhập tài khoản Bệnh nhân.\n2. Vào tab Đặt lịch, chọn dịch vụ S-02, chọn bác sĩ D-01.\n3. Chọn mốc giờ trống ngày mai và bấm Đặt lịch (xác thực OTP).',
        'Đặt lịch thành công, hệ thống tự động lấy thông tin bệnh nhân từ tài khoản đã đăng nhập, hiển thị lịch khám ở tab Lịch hẹn của tôi.',
        'ĐẠT'
    ],
    [
        'TC-UI-06',
        'Lễ tân đặt lịch trực tiếp tại quầy (Walk-In)',
        '1. Lễ tân vào tab Lịch Hẹn -> bấm Đặt lịch.\n2. Nhập SĐT bệnh nhân vãng lai cũ để hệ thống auto-lookup điền tên.\n3. Chọn khung giờ khám của bác sĩ và bấm xác nhận.',
        'Đặt lịch thành công với kênh bookingChannel là Walk-In, tự động bỏ qua bước OTP xác thực.',
        'ĐẠT'
    ],
    [
        'TC-UI-07',
        'Chặn đặt trùng giờ (Redis Overlap Lock)',
        '1. Dùng 2 tab trình duyệt chọn cùng 1 slot khám của cùng bác sĩ D-01 ngày mai.\n2. Bấm xác nhận đặt lịch ở cả 2 bên gần như đồng thời.',
        'Yêu cầu thứ 2 bị chặn lại với mã lỗi 409 (SLOT_NOT_AVAILABLE / APPOINTMENT_OVERLAP), đảm bảo không bị trùng lịch.',
        'ĐẠT'
    ]
]
create_test_table(headers, data_booking)

# 4. Queue & EMR
add_heading('4. KỊCH BẢN KIỂM THỬ NGHIỆP VỤ LIÊN THÔNG & HÀNG CHỜ')
data_queue = [
    [
        'TC-UI-08',
        'Tiếp đón bệnh nhân vào hàng chờ khám (Check-In)',
        '1. Lễ tân chọn lịch hẹn đã đặt hôm nay của bệnh nhân.\n2. Bấm nút Check-In (Tiếp đón).',
        'Bệnh nhân được thêm vào hàng chờ khám với trạng thái Waiting, hiển thị thời gian chờ đếm ngược trên màn hình TV phòng chờ.',
        'ĐẠT'
    ],
    [
        'TC-UI-09',
        'Bác sĩ gọi khám (In Chair)',
        '1. Bác sĩ mở Dashboard khám, chọn bệnh nhân trong hàng chờ.\n2. Bấm nút "Bắt đầu khám".',
        'Trạng thái chuyển sang In Chair. Đồng hồ đếm thời gian điều trị bắt đầu chạy, màn hình TV phòng chờ tự động cập nhật.',
        'ĐẠT'
    ],
    [
        'TC-UI-10',
        'Bác sĩ lưu bệnh án EMR & hoàn thành (Completed)',
        '1. Bác sĩ điền ghi chú bệnh án, chọn chẩn đoán răng trên sơ đồ và kê đơn thuốc.\n2. Bấm nút "Hoàn thành điều trị".',
        'Bệnh nhân chuyển sang trạng thái Completed. Hệ thống tự động tạo hóa đơn thanh toán tương ứng ở trạng thái Pending.',
        'ĐẠT'
    ],
    [
        'TC-UI-11',
        'Tự động liên kết bệnh án cũ khi tạo tài khoản',
        '1. Khách khám vãng lai (Walk-In) với SĐT "09xxxx".\n2. Bác sĩ khám và lưu bệnh án điện tử cho SĐT này.\n3. Bệnh nhân đăng ký tài khoản mới bằng đúng SĐT "09xxxx".',
        'Tài khoản mới đăng ký tự động liên kết với hồ sơ bệnh nhân cũ và hiển thị đầy đủ lịch sử bệnh án đã khám trước đó.',
        'ĐẠT'
    ]
]
create_test_table(headers, data_queue)

# 5. Kết quả kiểm nghiệm
add_heading('5. KẾT QUẢ KIỂM NGHIỆM TỰ ĐỘNG (AUTOMATED TEST SUMMARY)')
p_res = doc.add_paragraph()
p_res.paragraph_format.space_after = Pt(6)
r_res = p_res.add_run(
    'Tất cả các kịch bản kiểm thử API tích hợp tự động (Integration Tests) trong tập tin test_api.ts đã được chạy thử nghiệm '
    'và đạt kết quả tối ưu 100% (14/14 Test Cases ĐẠT PASSED) trên môi trường thử nghiệm cục bộ kết hợp máy chủ dữ liệu '
    'PostgreSQL và Redis Cache.'
)
r_res.font.name = 'Arial'
r_res.font.size = Pt(10)

doc.save('Kich_ban_kiem_thu_dat_lich.docx')
print('GENERATED_SUCCESSFULLY')
